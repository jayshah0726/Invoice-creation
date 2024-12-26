import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
import os
import random

def set_header_format(cell, text, bold=True, font_size=12):
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

def format_amount(amount):
    return f"Rs. {int(amount):,}"

def get_start_year(year_folder):
    return year_folder.split('-')[0]

def format_date(date):
    return pd.to_datetime(date).strftime('%d-%m-%Y')

def format_address(address):
    # Split address by commas
    parts = [part.strip() for part in address.split(',')]
    
    # Organize into roughly equal thirds
    total_parts = len(parts)
    third = max(1, total_parts // 3)
    
    line1 = ', '.join(parts[:third])
    line2 = ', '.join(parts[third:2*third])
    line3 = ', '.join(parts[2*third:])
    
    return f"{line1}\n{line2}\n{line3}"

def get_advisor_pan(advisor_name, data):
    """
    Find PAN for an advisor using flexible name matching across the entire dataset.
    Uses substring matching while checking all TEAM MEMBER entries.
    """
    # Handle empty or invalid inputs
    if not advisor_name or data.empty:
        return "PAN not found"
        
    # Split the search name into parts
    search_parts = advisor_name.upper().split()
    
    def names_match(row_name, search_name_parts):
        if pd.isna(row_name):
            return False
        
        row_parts = str(row_name).upper().split()
        
        # Check for exact match first
        if ' '.join(row_parts) == ' '.join(search_name_parts):
            return True
            
        # Then check for partial matches
        matches = []
        for search_part in search_name_parts:
            part_match = False
            for row_part in row_parts:
                if search_part in row_part or row_part in search_part:
                    part_match = True
                    break
            matches.append(part_match)
        
        return all(matches)
    
    # Create a mask for matching names in the entire TEAM MEMBER column
    matching_rows = data['TEAM MEMBER'].apply(lambda x: names_match(str(x), search_parts))
    
    # Get all matching rows from the complete dataset
    matches = data[matching_rows]
    
    # Debug: Print matching rows to verify
    print(f"Found {len(matches)} matches for {advisor_name}")
    if not matches.empty:
        print("Matching entries:")
        for idx, row in matches.iterrows():
            print(f"Name: {row['TEAM MEMBER']}, PAN: {row['PAN']}")
    
    # Check each match for valid PAN
    for _, row in matches.iterrows():
        if pd.notna(row['PAN']) and str(row['PAN']).strip() != "":
            return row['PAN']
            
    return "PAN not found"

# List of fonts to choose from
FONT_LIST = ['Calibri', 'Arial', 'Aptos Display', 'Cambria']

def set_advisor_font(doc, advisor_name):
    font_name = random.choice(FONT_LIST)
    doc.styles['Normal'].font.name = font_name
    print(f"Using font '{font_name}' for advisor: {advisor_name}")
    return font_name

def add_consolidated_invoice_to_doc(doc, invoice_date, advisor, address, total_data, year_folder, font_name, is_first_page=False, full_data=None):
    if not is_first_page:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    formatted_date = format_date(invoice_date)
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date: {formatted_date}")
    date_run.bold = True
    date_run.font.name = font_name
    
    if address:
        formatted_address = format_address(address)
        from_para = doc.add_paragraph(f"From:\n{advisor}\n{formatted_address}")
        for run in from_para.runs:
            run.font.name = font_name
    else:
        from_para = doc.add_paragraph(f"From:\n{advisor}")
        for run in from_para.runs:
            run.font.name = font_name

    if int(year_folder.split('-')[0]) <= 2020:
        from_address = """To:
ELIXIR WEALTH MANAGEMENT PVT. LTD.
OFFICE NO. 112, 1ST FLOOR,
FORTUNE GEE BEE COMPLEX,
VAPI DAMAN MAIN ROAD
SOMNATH, DAMAN – 396210"""
    else:  
        from_address = """To:
DIPAN MEHTA COMMODITIES PVT. LTD
58 MITTAL CHAMBERS,
228, NARIMAN POINT
MUMBAI 400 021"""
    
    doc.add_paragraph(from_address)

    instruction_paragraph = doc.add_paragraph()
    instruction_run = instruction_paragraph.add_run('INSTRUCTION NOTE')
    instruction_run.bold = True
    instruction_paragraph.alignment = 1

    table = doc.add_table(rows=1, cols=0)
    table.style = 'Table Grid'
    table.autofit = False

    table.add_column(Inches(0.5))    # Sr. No
    table.add_column(Inches(4.7))    # Particulars
    table.add_column(Inches(1.3))    # Amount

    headers = ['Sr. No', 'Particulars', 'Amount']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_header_format(cell, header)

    row_cells = table.add_row().cells
    total_amount = total_data['TOTAL AMOUNT'].sum()
    
    row_cells[0].text = '1'
    start_year = get_start_year(year_folder)
    row_cells[1].text = f'Being amount payable for the services rendered as per the engagement letter dated 1st April {start_year}'
    row_cells[2].text = format_amount(total_amount)

    total_row = table.add_row().cells
    total_row[1].text = 'Total'
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = format_amount(total_amount)
    total_row[2].paragraphs[0].runs[0].bold = True

    doc.add_paragraph("""
Terms & Conditions:
1) E&OE
2) Please arrange to make the payments to payees as detailed in Annexure "A"
3) Please arrange to pay the said amount within 10 working days.
4) You are requested to pay total amount Net of TDS and arrange to send TDS certificate.
    """)

    advisor_pan = get_advisor_pan(advisor, full_data)  # Pass the full dataset here
    regards_para = doc.add_paragraph("Regards,")
    
    for _ in range(3):
        doc.add_paragraph()
    
    doc.add_paragraph(f"Name: {advisor}")
    doc.add_paragraph(f"PAN: {advisor_pan}")

def add_consolidated_annexure_a_to_doc(doc, payee_data, font_name):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    annexure_header = doc.add_paragraph()
    header_run = annexure_header.add_run('Annexure A')
    header_run.bold = True
    header_run.font.name = font_name
    
    doc.add_paragraph('Details of payees to whom payment is to be made')

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['NAME', 'PAN', 'Fees', 'OUT OF POCKET', 'TOTAL']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_header_format(cell, header)

    for _, row in payee_data.iterrows():
        new_row = table.add_row()
        cells = new_row.cells
        
        cells[0].text = str(row['TEAM MEMBER'])
        cells[1].text = str(row['PAN'])
        cells[2].text = format_amount(row['PROFESSIONAL FEES'])
        cells[3].text = format_amount(row['OUT OF POCKET'])
        cells[4].text = format_amount(row['TOTAL AMOUNT'])

    total_row = table.add_row().cells
    total_row[0].text = 'Total'
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = ''
    total_row[2].text = ''
    total_row[3].text = ''
    total_amount = payee_data['TOTAL AMOUNT'].sum()
    total_row[4].text = format_amount(total_amount)
    total_row[4].paragraphs[0].runs[0].bold = True

def create_consolidated_invoices_for_advisor(grouped_data, address_data, year_folder, trading_advisor, full_data):
    trading_advisor = trading_advisor.strip()
    doc = Document()
    
    font_name = set_advisor_font(doc, trading_advisor)

    for idx, (invoice_date, date_group) in enumerate(grouped_data.groupby('Invoice Date')):
        first_row = date_group.iloc[0].copy()
    
        address_data['NAME'] = address_data['NAME'].str.strip()
        first_row['TRADING ADVISOR'] = first_row['TRADING ADVISOR'].strip()
        
        address = address_data[address_data['NAME'] == first_row['TRADING ADVISOR']]['Address'].values
        
        if len(address) == 0:
            trading_advisor = first_row['TRADING ADVISOR']
            trading_advisor_parts = trading_advisor.split()
            if len(trading_advisor_parts) >= 2:
                first_name = trading_advisor_parts[0]
                last_name = trading_advisor_parts[-1]
                address = address_data[(address_data['TRADING ADVISOR'].str.contains(first_name)) & (address_data['TRADING ADVISOR'].str.contains(last_name))]['Address'].values
        
        address = address[0] if len(address) > 0 else 'Address not found'

        add_consolidated_invoice_to_doc(
            doc,
            invoice_date=invoice_date,
            advisor=first_row['TRADING ADVISOR'],
            address=address,
            total_data=date_group,
            year_folder=year_folder,
            font_name=font_name,
            is_first_page=(idx == 0),
            full_data=full_data
        )

        add_consolidated_annexure_a_to_doc(doc, date_group, font_name)

    payee_folder = os.path.join(year_folder)
    os.makedirs(payee_folder, exist_ok=True)
    file_name = os.path.join(payee_folder, f"Combined_Invoice_{trading_advisor.replace(' ', '_')}.docx")
    doc.save(file_name)
    print(f"Consolidated Invoice saved: {file_name}")

# Main execution logic
file_path = 'Payment of Professional Fees 2.xlsx'
address_file_path = 'DMC-Master-Data-1.csv' 
sheet_names = ['2022-23']

address_data = pd.read_csv(address_file_path)

for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')
    data = pd.read_excel(file_path, sheet_name='REVISED PAYMENTS TO TRADERS')
    grouped_data_by_advisor = data.groupby('TRADING ADVISOR')
    
    print(f"\nProcessing year: {year_folder}")
    print(f"Total number of invoices to be generated: {len(grouped_data_by_advisor)}")
    
    for trading_advisor, group in grouped_data_by_advisor:
        create_consolidated_invoices_for_advisor(group, address_data, year_folder, trading_advisor, data)
    
    print(f"Completed processing {len(grouped_data_by_advisor)} invoices for {year_folder}")