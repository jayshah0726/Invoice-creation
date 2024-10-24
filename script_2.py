import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import os

# Function to format the table headers
def set_header_format(cell, text, bold=True, font_size=12):
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

# Function to format currency amounts
def format_amount(amount):
    return f"Rs. {int(amount):,}"  # Remove decimals and add thousands separator

# Function to calculate the start year from the fiscal year folder name
def get_start_year(year_folder):
    return year_folder.split('-')[0]

# Helper function to format the date in DD-MM-YYYY
def format_date(date):
    return pd.to_datetime(date).strftime('%d-%m-%Y')

# Function to format address into four lines
def format_address(address):
    # Split address by commas
    parts = [part.strip() for part in address.split(',')]
    
    # Ensure the city and pin code appear on the fourth line
    city_and_pin = parts[-1] if len(parts) > 0 else ''  # Get the last part (city and pin code)
    remaining_parts = parts[:-1]  # All other parts of the address
    
    # Organize the remaining parts into roughly equal thirds
    total_parts = len(remaining_parts)
    third = max(1, total_parts // 3)  # Divide into three equal parts, ensure at least 1 part per line
    
    line1 = ', '.join(remaining_parts[:third])
    line2 = ', '.join(remaining_parts[third:2*third])
    line3 = ', '.join(remaining_parts[2*third:])
    
    # Format the address with the first three lines and then the city and pin code in the fourth line
    return f"{line1}\n{line2}\n{line3}\n{city_and_pin}"

# Function to add an invoice for a unique date
def add_invoice_to_doc(doc, invoice_date, advisor, pan, address, professional_fees, out_of_pocket, total_amount, year_folder, is_first_page=False):
    # Add page break for all pages except the first one
    if not is_first_page:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Add content to the document for each unique payment date
    formatted_date = format_date(invoice_date)
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date: {formatted_date}")
    date_run.bold = True
    
    
    if address:
        formatted_address = format_address(address)
        doc.add_paragraph(f"From:\n{advisor}\n{formatted_address}")
    else:
        doc.add_paragraph(f"From:\n{advisor}")

    doc.add_paragraph("""
To:
ELIXIR WEALTH MANAGEMENT PVT. LTD.
58 MITTAL CHAMBERS,
228, NARIMAN POINT
MUMBAI 400 021
    """)

    instruction_paragraph = doc.add_paragraph()
    instruction_run = instruction_paragraph.add_run('INSTRUCTION NOTE')
    instruction_run.bold = True
    instruction_paragraph.alignment = 1  # 1 represents center alignment

    # Table for service details
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Format table headers
    hdr_cells = table.rows[0].cells
    set_header_format(hdr_cells[0], 'Sr. No', bold=True, font_size=12)
    set_header_format(hdr_cells[1], 'Particulars', bold=True, font_size=12)
    set_header_format(hdr_cells[2], 'Amount', bold=True, font_size=12)

    # Add service row details
    row_cells = table.rows[1].cells
    row_cells[0].text = '1'
    start_year = get_start_year(year_folder)
    row_cells[1].text = f'Being amount payable for the services rendered as per the engagement letter dated 1st April {start_year}'
    row_cells[2].text = format_amount(total_amount)

    # Add total row
    total_row = table.add_row().cells
    total_row[1].text = 'Total'
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = format_amount(total_amount)
    total_row[2].paragraphs[0].runs[0].bold = True

    # Add terms and conditions
    doc.add_paragraph("""
Terms & Conditions:
1) E&OE
2) Please arrange to make the payments to payees as detailed in Annexure "A"
3) Please arrange to pay the said amount within 10 working days.
4) You are requested to pay total amount Net of TDS and arrange to send TDS certificate.
    """)

    doc.add_paragraph(f"Regards,\n\nName: {advisor}\nPAN: {pan}")

# Function to add Annexure A to the document
def add_annexure_a_to_doc(doc, payee, pan, professional_fees, out_of_pocket, total_amount):
    # Add Annexure A header
    doc.add_paragraph().add_run('Annexure A').bold = True
    doc.add_paragraph('Details of payees to whom payment is to be made')

    # Table for payee details
    annexure_table = doc.add_table(rows=2, cols=5)
    annexure_table.style = 'Table Grid'
    
    # Set headers
    headers = ['NAME', 'PAN', 'Fees', 'OUT OF POCKET', 'TOTAL']
    for i, header in enumerate(headers):
        cell = annexure_table.rows[0].cells[i]
        set_header_format(cell, header)

    # Add data with formatted amounts
    annexure_row = annexure_table.rows[1].cells
    annexure_row[0].text = payee
    annexure_row[1].text = pan
    annexure_row[2].text = format_amount(professional_fees)
    annexure_row[3].text = format_amount(out_of_pocket)
    annexure_row[4].text = format_amount(total_amount)

# Function to create combined invoices for each trading advisor
def create_combined_invoices_for_advisor(grouped_data, address_data, year_folder, trading_advisor):
    trading_advisor = trading_advisor.strip()
    doc = Document()
    
    # Process each unique payment date
    for idx, (invoice_date, group) in enumerate(grouped_data.groupby('Invoice Date')):
        for _, row in group.iterrows():
            # Get address
            address = address_data[address_data['Name'] == row['TRADING ADVISOR']]['Address'].values
            address = address[0] if len(address) > 0 else 'Address not found'

            # Add invoice with proper page break handling
            add_invoice_to_doc(
                doc,
                invoice_date=invoice_date,
                advisor=row['TRADING ADVISOR'],
                pan=row['PAN'],
                address=address,
                professional_fees=row['PROFESSIONAL FEES'],
                out_of_pocket=row['OUT OF POCKET'],
                total_amount=row['TOTAL AMOUNT'],
                year_folder=year_folder,
                is_first_page=(idx == 0)  # Only first invoice doesn't need a page break
            )

            add_annexure_a_to_doc(
                doc,
                payee=row['PAYEE'],
                pan=row['PAN'],
                professional_fees=row['PROFESSIONAL FEES'],
                out_of_pocket=row['OUT OF POCKET'],
                total_amount=row['TOTAL AMOUNT']
            )

    # Save the document
    payee_folder = os.path.join(year_folder, trading_advisor.replace(' ', '_'))
    os.makedirs(payee_folder, exist_ok=True)
    file_name = os.path.join(payee_folder, f"Combined_Invoice_{trading_advisor.replace(' ', '_')}.docx")
    doc.save(file_name)
    print(f"Combined Invoice saved: {file_name}")

# Main execution logic
file_path = 'Latest EW - Details of Professional Fees Paid for last 7 years.xlsx'
address_file_path = 'EW Master.csv' 
sheet_names = ['2017-18', '2018-19', '2019-20', '2020-21', '2021-22', '2023-24']

address_data = pd.read_csv(address_file_path)

for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')
    data = pd.read_excel(file_path, sheet_name=sheet)
    grouped_data_by_advisor = data.groupby('TRADING ADVISOR')
    
    for trading_advisor, group in grouped_data_by_advisor:
        create_combined_invoices_for_advisor(group, address_data, year_folder, trading_advisor)
