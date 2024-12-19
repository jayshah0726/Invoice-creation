from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import textwrap
import os

master_file_path = 'Payment of Professional Fees 2.xlsx'
address_file_path = 'DMC-Master-Data-1.csv' 
sheet_names = ['REVISED PAYMENTS TO TRADERS']

# '2022-23'

address_data = pd.read_csv(address_file_path)

def format_address(address, max_line_length=40):
    """
    Formats the given address into a string with up to 4 lines.
    
    Parameters:
    - address (str): The full address to be formatted.
    - max_line_length (int): Maximum number of characters per line. Default is 40.
    
    Returns:
    - str: The formatted address with up to 4 lines.
    """
    # Wrap the address into lines of specified width
    address_lines = textwrap.wrap(address, width=max_line_length)

    # Ensure there are exactly 4 lines by padding with empty strings if necessary
    while len(address_lines) < 4:
        address_lines.append("")

    # Join the first 4 lines with newline characters
    formatted_address = "\n".join(address_lines[:4])
    
    return formatted_address

def set_font(paragraph, font_name="Roboto", font_size=11):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = None  # Set text color to black

        # Ensures font name works with Word documents
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

# Set font for all text in a table
def set_table_font(table, font_name="Roboto", font_size=11):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name=font_name, font_size=font_size)

for sheet in sheet_names:
    year_folder = "2022-23"
    master_data = pd.read_csv(address_file_path)
    data = pd.read_excel(master_file_path, sheet_name=sheet)
    trading_advisor_list = list(data['TRADING ADVISOR'].drop_duplicates().dropna())
    filtered_data = data[['TRADING ADVISOR', 'TEAM MEMBER', 'PAN']]
    grouped_data_by_advisor = filtered_data.groupby('TRADING ADVISOR')
    ta_payee_mapping = grouped_data_by_advisor.agg({
        "PAN": list,
        "TEAM MEMBER": list
    }).reset_index()

    for current_trading_advisor in trading_advisor_list:
        current_trading_advisor = current_trading_advisor.strip()
        print(current_trading_advisor)
        if current_trading_advisor == "MITESH DOSHI":
            current_trading_advisor = "MITESH JAYANTIBHAI DOSHI"
        elif current_trading_advisor == "MITUL MORABIA":
            current_trading_advisor = "MITUL MOHANLAL MORABIYA"

        # Address lookup
        current_trading_advisor_new = master_data[(master_data['NAME'] == current_trading_advisor)]
        current_trading_advisor_address = current_trading_advisor_new['Address'].iloc[0]
        print(current_trading_advisor_address)
        address_text = format_address(current_trading_advisor_address)

        # Initialize the document
        doc = Document()

        # Title
        title = doc.add_paragraph("ENGAGEMENT LETTER\n")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(title, font_name="Roboto", font_size=11)

        # Date
        date = doc.add_paragraph(f"April 01, {year_folder}", style='Normal')
        set_font(date, font_name="Roboto", font_size=11)

        # "To" section
        to_para = doc.add_paragraph()
        to_para.add_run(f"To,\n{current_trading_advisor},\n{address_text}")
        set_font(to_para, font_name="Roboto", font_size=11)

        # Subject line
        subject_para = doc.add_paragraph("Subject: Appointment as Trading Advisor")
        set_font(subject_para, font_name="Roboto", font_size=11)

        # Introduction
        intro = doc.add_paragraph(
            "ELIXIR EQUITIES PVT. LTD (hereinafter referred to as “the Company”) is a company incorporated under the erstwhile provisions of the Companies Act, 1956 carrying on the business of securities trading across various market segments."
        )
        set_font(intro, font_name="Roboto", font_size=11)

        # Content paragraph
        content = doc.add_paragraph(
            "Considering your expertise in the subject matter, the Company is desirous of appointing your goodself as a trading advisor (hereinafter referred to as “TA”) with respect to its trading operations and undertake trading activities on its behalf including but not limited to trading, jobbing, arbitrage, hedging etc."
        )
        set_font(content, font_name="Roboto", font_size=11)

        # Terms and Conditions section
        tnc = doc.add_paragraph("The terms and conditions for your appointment would be as under:")
        set_font(tnc, font_name="Roboto", font_size=12)

        terms_and_conditions = [
            "The Company shall provide all the necessary infrastructure to you, including trading terminals of the Bombay Stock Exchange and the National Stock Exchange, trading/algo software, charting software, Wi-Fi, internet, web access, television, furniture and fixtures, electricity, water, air conditioning, telephone/s lines, etc.",
            "We understand that you would be assisted by your team members (hereinafter referred to as 'the team/team members'), the details of which are as per Annexure 'A'.",
            "You and your team shall be permitted to use the facilities described in para 1 above. You and team shall devote your skill, time, ability and attention to conducting trading/jobbing/arbitrage/hedging transactions/strategies on the exchange/s in the best interest of the Company. You and your team shall be responsible for all trading-related activities of the team.",
            "You shall execute/instruct to execute all the transactions in the Unique Client Code of the Company and only for the Company. The Company shall also provide the requisite funds to enable you to undertake transactions on the exchange/s through a SEBI registered Broker in the Company’s client code.",
            "You and your team agree to keep an interest-free security deposit, as may be mutually agreed from time to time, the proceeds of which will be utilized by the Company at its sole discretion.",
            "You and your team agree to carry on said business in complete confidentiality and shall not divulge trading positions, strategies, etc., or any other information related to the said business to anyone whatsoever.",
            "You shall charge fees for the services rendered by you and your team and shall periodically provide necessary instructions for disbursal of the same. The Company will make payment to you and your team as detailed in the instruction note after appropriate deduction of tax at source as per the provisions of the Income-tax Act, 1961.",
            "As mentioned earlier, you may also, if agreed upon by the Company, appoint/employ other persons referred to as your team to assist you in your trading activity. Your team agrees to abide by the terms and conditions as set out in this letter and confirm the same. New team members may be introduced from time to time at your sole discretion. Necessary intimation will be sent by you to the Company on introduction or removal of any team member.",
            "You and your team shall implement the terms of this agreement in good faith and due diligence in the best interest of the Company.",
            "Nothing in this arrangement shall constitute or be deemed to constitute a partnership, relationship of principal and agent or employer and employee between any of the parties, and none of them shall have any authority to bind any of the other parties in any way except for the purposes of the business of the Company.",
            "You or your team members shall not undertake any personal trading in the Unique Client Code of the Company under any circumstances. This shall be construed as 'unauthorized' trading activity and liable for damages by the Company.",
            "You and your team shall hereby comply with all the applicable rules and regulations, without limitation to, SEBI (Prohibition of Fraudulent and Unfair Trade Practices Relating to Securities Market) Regulations, Exchange guidelines and regulations, and any amendments and changes thereto, or any other act/s, and any such guidelines as may be prescribed from time to time by the Company.",
            "This arrangement shall be for one year from April 01, 2023, but the Company reserves the right to terminate at any time without giving any prior notice or modify any terms and conditions of this letter from time to time as may be deemed necessary by the Company."
        ]

        for term in terms_and_conditions:
            paragraph = doc.add_paragraph(term+"\n", style='List Number')
            set_font(paragraph, font_name="Roboto", font_size=11)

        # Signature section
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.style = 'Table Grid'
        sign_table.autofit = True

        sign_table.cell(0, 0).text = "For ELIXIR EQUITIES PVT. LTD,\n\n\n\nDipan Mehta\nDirector"
        sign_table.cell(0, 1).text = f"I Accept\n\n\n\nName: {current_trading_advisor} \nPAN : {current_trading_advisor_new['PAN'].iloc[0]}"

        set_table_font(sign_table, font_name="Roboto", font_size=11)

        # Annexure A on a new page
        doc.add_page_break()
        annexure = doc.add_paragraph("Annexure A\n")
        annexure.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(annexure, font_name="Roboto", font_size=11)
        members = doc.add_paragraph(
            "The following is the list of team members of the TA presently working with him which may change from time to time as per commercial prudence of the TA.\n"
        )
        set_font(members, font_name="Roboto", font_size=11)
        
        filtered_ta_payee_mapping = ta_payee_mapping[ta_payee_mapping['TRADING ADVISOR'] == current_trading_advisor]

        for index, each_row in filtered_ta_payee_mapping.iterrows():
            trading_advisor = each_row['TRADING ADVISOR']
            pan_list = each_row["PAN"]
            payee_list = each_row['TEAM MEMBER']
            pan_payee_dict = {}

            # Populate pan_payee_dict with valid entries
            for i in range(len(pan_list)):
                if payee_list[i] not in pan_payee_dict and payee_list[i] != trading_advisor:
                    pan_payee_dict[payee_list[i]] = pan_list[i]

            # Create the table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Add headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'PAYEE'
            hdr_cells[1].text = 'PAN'
            hdr_cells[2].text = 'SIGNATURE'

            # If pan_payee_dict is not empty, add rows for the payees
            if len(pan_payee_dict) > 0:
                for key, value in pan_payee_dict.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
            else:
                # If pan_payee_dict is empty, add two empty rows
                for _ in range(2):
                    dummy_row = table.add_row().cells
                    dummy_row[0].text = ''
                    dummy_row[1].text = ''
                    dummy_row[2].text = ''

            # Apply font style to the entire table
            set_table_font(table, font_name="Roboto", font_size=12)
        
        os.makedirs(f"EEPL-{year_folder}", exist_ok=True)
        doc.save(f"EEPL-{year_folder}/Engagement_Letter_{current_trading_advisor}_{year_folder}.docx")
