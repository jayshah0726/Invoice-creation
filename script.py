import pandas as pd
from docx import Document
from docx.shared import Pt
import os

# Function to format the table headers
def set_header_format(cell, text, bold=True, font_size=12):
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

# Function to calculate the start year from the fiscal year folder name (e.g., '2017-18' -> '2017')
def get_start_year(year_folder):
    return year_folder.split('-')[0]

# Function to add an invoice for a unique date
def add_invoice_to_doc(doc, payment_date, payee, pan, professional_fees, out_of_pocket, total_amount, year_folder):
    # Add content to the document for each unique payment date
    doc.add_paragraph(f"Date: {payment_date}")
    doc.add_paragraph(f"From: {payee}\n")

    doc.add_paragraph("""
To:
ELIXIR WEALTH MANAGEMENT PVT. LTD.
58 MITTAL CHAMBERS,
228, NARIMAN POINT
MUMBAI 400 021
    """)

    # Add 'Instruction Note' header (non-blue, bold)
    doc.add_paragraph().add_run('INSTRUCTION NOTE').bold = True

    # Table for service details (3 columns)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Format table headers (Sr. No, Particulars, Amount)
    hdr_cells = table.rows[0].cells
    set_header_format(hdr_cells[0], 'Sr. No', bold=True, font_size=12)
    set_header_format(hdr_cells[1], 'Particulars', bold=True, font_size=12)
    set_header_format(hdr_cells[2], 'Amount (Rs.)', bold=True, font_size=12)

    # Add service row details
    row_cells = table.rows[1].cells
    row_cells[0].text = '1'
    
    # Extract the correct start year (e.g., 2017 for 2017-18)
    start_year = get_start_year(year_folder)
    row_cells[1].text = f'Being amount payable for the services rendered as per the engagement letter dated 1st April {start_year}'
    row_cells[2].text = f'{total_amount}'

    # Add a new row for "Total"
    total_row = table.add_row().cells
    total_row[1].text = 'Total'
    total_row[1].paragraphs[0].runs[0].bold = True  # Make "Total" bold
    total_row[2].text = f'{total_amount}'
    total_row[2].paragraphs[0].runs[0].bold = True  # Make the amount bold as well

    # Add terms and conditions
    doc.add_paragraph("""
Terms & Conditions:
1) E&OE
2) Please arrange to make the payments to payees as detailed in Annexure “A”
3) Please arrange to pay the said amount within 10 working days.
4) You are requested to pay total amount Net of TDS and arrange to send TDS certificate.
    """)

    doc.add_paragraph(f"Regards,\n\nName: {payee}\nPAN: {pan}")

# Function to add Annexure A to the document
def add_annexure_a_to_doc(doc, payee, pan, professional_fees, out_of_pocket, total_amount):
    # Add Annexure A header (non-blue, bold)
    doc.add_paragraph().add_run('Annexure A').bold = True
    doc.add_paragraph('Details of payees to whom payment is to be made')

    # Table for payee details
    annexure_table = doc.add_table(rows=2, cols=5)
    annexure_table.style = 'Table Grid'
    annexure_hdr = annexure_table.rows[0].cells
    annexure_hdr[0].text = 'NAME'
    annexure_hdr[1].text = 'PAN'
    annexure_hdr[2].text = 'Fees'
    annexure_hdr[3].text = 'OUT OF POCKET'
    annexure_hdr[4].text = 'TOTAL'

    annexure_row = annexure_table.rows[1].cells
    annexure_row[0].text = payee
    annexure_row[1].text = pan
    annexure_row[2].text = str(professional_fees)
    annexure_row[3].text = str(out_of_pocket)
    annexure_row[4].text = str(total_amount)

# Function to create invoices for unique payment dates under the same document for each trading advisor
def create_combined_invoices_for_advisor(grouped_data, year_folder, trading_advisor):
    # Create a new document
    doc = Document()

    # Step 1: Process each unique payment date for the trading advisor
    for payment_date, group in grouped_data.groupby('Payment Date'):
        for _, row in group.iterrows():
            # Add invoice details to the document
            add_invoice_to_doc(
                doc,
                payment_date=payment_date,
                payee=row['PAYEE'],
                pan=row['PAN'],
                professional_fees=row['PROFESSIONAL FEES'],
                out_of_pocket=row['OUT OF POCKET'],
                total_amount=row['TOTAL AMOUNT'],
                year_folder=year_folder
            )

    # Step 2: Add Annexure A details at the end
    # Assume that the annexure is the same for each payment
    last_row = group.iloc[-1]  # Use the last row for annexure details
    add_annexure_a_to_doc(
        doc,
        payee=last_row['PAYEE'],
        pan=last_row['PAN'],
        professional_fees=last_row['PROFESSIONAL FEES'],
        out_of_pocket=last_row['OUT OF POCKET'],
        total_amount=last_row['TOTAL AMOUNT']
    )

    # Step 3: Save the document after all invoices are added
    payee_folder = os.path.join(year_folder, trading_advisor)
    os.makedirs(payee_folder, exist_ok=True)

    file_name = os.path.join(payee_folder, f"Combined_Invoice_{trading_advisor}.docx")
    doc.save(file_name)
    print(f"Combined Invoice saved: {file_name}")

# Main logic to read each year and generate invoices
file_path = 'Latest EW - Details of Professional Fees Paid for last 7 years.xlsx'
sheet_names = ['2017-18', '2018-19', '2019-20', '2020-21', '2021-22', '2022-23', '2023-24']

# Step 1: Loop through each sheet (year)
for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')  # Replace / to avoid issues with folder names

    # Read the data for the current year
    data = pd.read_excel(file_path, sheet_name=sheet)

    # Step 2: Group data by TRADING ADVISOR
    grouped_data_by_advisor = data.groupby('TRADING ADVISOR')

    # Step 3: Loop through grouped data for each trading advisor
    for trading_advisor, group in grouped_data_by_advisor:
        create_combined_invoices_for_advisor(group, year_folder, trading_advisor)