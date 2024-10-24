import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
import os

# Function to format the table headers
def set_header_format(cell, text, bold=True, font_size=12):
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)

# Function to format currency amounts
def format_amount(amount):
    return f"Rs. {int(amount):,}"

def get_start_year(year_folder):
    return year_folder.split('-')[0]

def format_date(date):
    return pd.to_datetime(date).strftime('%d-%m-%Y')

def format_address(address):
    parts = [part.strip() for part in address.split(',')]
    city_and_pin = parts[-1] if len(parts) > 0 else ''
    remaining_parts = parts[:-1]
    
    total_parts = len(remaining_parts)
    third = max(1, total_parts // 3)
    
    line1 = ', '.join(remaining_parts[:third])
    line2 = ', '.join(remaining_parts[third:2*third])
    line3 = ', '.join(remaining_parts[2*third:])
    
    return f"{line1}\n{line2}\n{line3}\n{city_and_pin}"

def add_consolidated_invoice_to_doc(doc, invoice_date, advisor, address, total_data, year_folder, is_first_page=False):
    # Always add a page break unless it's the first page
    if not is_first_page:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    formatted_date = format_date(invoice_date)
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date: {formatted_date}")
    date_run.bold = True
    
    if address:
        formatted_address = format_address(address)
        doc.add_paragraph(f"From:\n{advisor}\n{formatted_address}")
    else:
        doc.add_paragraph(f"From:\n{advisor}")

    doc.add_paragraph("""
To:
ELIXIR WEALTH MANAGEMENT PVT. LTD.
58 MITTAL CHAMBERS,
228, NARIMAN POINT
MUMBAI 400 021
    """)

    instruction_paragraph = doc.add_paragraph()
    instruction_run = instruction_paragraph.add_run('INSTRUCTION NOTE')
    instruction_run.bold = True
    instruction_paragraph.alignment = 1

    # Table for service details with consolidated amount
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Set the width for each column
    table.autofit = True
    widths = (Inches(0.4), Inches(4.5), Inches(1.5))  # Updated widths for columns
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Formatting header cells
    hdr_cells = table.rows[0].cells
    set_header_format(hdr_cells[0], 'Sr. No', bold=True, font_size=12)
    set_header_format(hdr_cells[1], 'Particulars', bold=True, font_size=12)
    set_header_format(hdr_cells[2], 'Amount', bold=True, font_size=12)

    # Add data to the table
    total_amount = total_data['TOTAL AMOUNT'].sum()
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '1'
    start_year = get_start_year(year_folder)
    row_cells[1].text = f'Being amount payable for the services rendered as per the engagement letter dated 1st April {start_year}'
    row_cells[2].text = format_amount(total_amount)

    # Add total row
    total_row = table.add_row().cells
    total_row[1].text = 'Total'
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = format_amount(total_amount)
    total_row[2].paragraphs[0].runs[0].bold = True

    # Add terms and conditions
    doc.add_paragraph("""
Terms & Conditions:
1) E&OE
2) Please arrange to make the payments to payees as detailed in Annexure "A"
3) Please arrange to pay the said amount within 10 working days.
4) You are requested to pay total amount Net of TDS and arrange to send TDS certificate.
    """)

    # Get PAN from the first row of data (assuming same PAN for advisor)
    advisor_pan = total_data.iloc[0]['PAN'] if not pd.isna(total_data.iloc[0]['PAN']) else "PAN not found"
    
    # Add regards with proper spacing
    regards_para = doc.add_paragraph("Regards,")
    
    # Add 4 empty lines
    for _ in range(4):
        doc.add_paragraph()
    
    # Add name and PAN
    doc.add_paragraph(f"Name: {advisor}")
    doc.add_paragraph(f"PAN: {advisor_pan}")


def add_consolidated_annexure_a_to_doc(doc, payee_data):
    # Always start annexure on a new page
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Add Annexure header
    annexure_header = doc.add_paragraph()
    header_run = annexure_header.add_run('Annexure A')
    header_run.bold = True
    
    doc.add_paragraph('Details of payees to whom payment is to be made')

    # Create table with exact number of rows needed
    table = doc.add_table(rows=1, cols=5)  # Start with header row only
    table.style = 'Table Grid'
    
    # Set headers
    headers = ['NAME', 'PAN', 'Fees', 'OUT OF POCKET', 'TOTAL']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_header_format(cell, header)

    # Add rows one by one for each payee
    for _, row in payee_data.iterrows():
        # Add a new row
        new_row = table.add_row()
        cells = new_row.cells
        
        # Fill in the data
        cells[0].text = str(row['PAYEE'])
        cells[1].text = str(row['PAN'])
        cells[2].text = format_amount(row['PROFESSIONAL FEES'])
        cells[3].text = format_amount(row['OUT OF POCKET'])
        cells[4].text = format_amount(row['TOTAL AMOUNT'])

def create_consolidated_invoices_for_advisor(grouped_data, address_data, year_folder, trading_advisor):
    trading_advisor = trading_advisor.strip()
    doc = Document()
    
    # Group by date to consolidate invoices
    for idx, (invoice_date, date_group) in enumerate(grouped_data.groupby('Invoice Date')):
        # Get address from first row (assuming same address for advisor)
        first_row = date_group.iloc[0]
        address = address_data[address_data['Name'] == first_row['TRADING ADVISOR']]['Address'].values
        address = address[0] if len(address) > 0 else 'Address not found'

        # Add consolidated invoice
        add_consolidated_invoice_to_doc(
            doc,
            invoice_date=invoice_date,
            advisor=first_row['TRADING ADVISOR'],
            address=address,
            total_data=date_group,
            year_folder=year_folder,
            is_first_page=(idx == 0)
        )

        # Add consolidated Annexure A with all payees for this date
        add_consolidated_annexure_a_to_doc(doc, date_group)

    # Save the document
    payee_folder = os.path.join(year_folder, trading_advisor.replace(' ', '_'))
    os.makedirs(payee_folder, exist_ok=True)
    file_name = os.path.join(payee_folder, f"Combined_Invoice_{trading_advisor.replace(' ', '_')}.docx")
    doc.save(file_name)
    print(f"Consolidated Invoice saved: {file_name}")

# Main execution logic
file_path = 'Latest EW - Details of Professional Fees Paid for last 7 years.xlsx'
address_file_path = 'EW Master.csv' 
sheet_names = ['2017-18', '2018-19', '2019-20', '2020-21', '2021-22', '2023-24']

address_data = pd.read_csv(address_file_path)

for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')
    data = pd.read_excel(file_path, sheet_name=sheet)
    grouped_data_by_advisor = data.groupby('TRADING ADVISOR')
    
    for trading_advisor, group in grouped_data_by_advisor:
        create_consolidated_invoices_for_advisor(group, address_data, year_folder, trading_advisor)