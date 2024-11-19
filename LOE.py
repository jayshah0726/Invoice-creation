from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import os

file_path = 'EW - Details of Professional Fees Paid for last 7 years 1.xlsx'
address_file_path = 'EW Master 1(Master Data).csv' 
sheet_names = ['2018-19']

address_data = pd.read_csv(address_file_path)

def set_font(paragraph, font_name="Roboto", font_size=11):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = None  # Set text color to black

        # Ensures font name works with Word documents
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

# Set font for all text in a table
def set_table_font(table, font_name="Roboto", font_size=11):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name=font_name, font_size=font_size)

for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')
    master_data = pd.read_csv('EW Master 1(Master Data).csv')
    data = pd.read_excel('EW - Details of Professional Fees Paid for last 7 years.xlsx', sheet_name=sheet)
    trading_advisor_list = list(data['TRADING ADVISOR'].drop_duplicates().dropna())
    filtered_data = data[['TRADING ADVISOR', 'PAYEE', 'PAN']]
    grouped_data_by_advisor = filtered_data.groupby('TRADING ADVISOR')
    ta_payee_mapping = grouped_data_by_advisor.agg({
        "PAN": list,
        "PAYEE": list
    }).reset_index()

    for current_trading_advisor in trading_advisor_list:
        current_trading_advisor = current_trading_advisor.strip()
        if current_trading_advisor == "MITESH DOSHI":
            current_trading_advisor = "MITESH JAYANTIBHAI DOSHI"
        elif current_trading_advisor == "MITUL MORABIA":
            current_trading_advisor = "MITUL MOHANLAL MORABIYA"

        # Address lookup
        current_trading_advisor_new = master_data[master_data['Name'] == current_trading_advisor]
        current_trading_advisor_address = current_trading_advisor_new['Address'].iloc[0].split(", ")
        address_text = "\n".join(current_trading_advisor_address[:4])

        # Initialize the document
        doc = Document()

        # Title
        title = doc.add_paragraph("ENGAGEMENT LETTER", style='Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(title, font_name="Roboto", font_size=11)

        # Date
        date = doc.add_paragraph(f"April 01, {sheet.split('-')[0]}", style='Normal')
        set_font(date, font_name="Roboto", font_size=11)
        doc.add_paragraph()  # Blank row

        # Add "To" section
        to_para = doc.add_paragraph()
        to_para.add_run(f"To,\n{current_trading_advisor},\n{address_text}\n")
        set_font(to_para, font_name="Roboto", font_size=11)
        doc.add_paragraph()  # Blank row

        # Subject line
        subject_para = doc.add_paragraph()
        subject_para.add_run("Subject: Appointment as Trading Advisor\n\n").bold = False
        set_font(subject_para, font_name="Roboto", font_size=11)
        doc.add_paragraph()  # Blank row

        # Introduction
        intro = doc.add_paragraph(
            "Elixir Equities Pvt Ltd (hereinafter referred to as “the Company”) is a company incorporated under the erstwhile provisions of the Companies Act, 1956 carrying on the business of securities trading across various market segments."
        )
        set_font(intro, font_name="Roboto", font_size=11)
        doc.add_paragraph()  # Blank row

        # Content paragraph
        content = doc.add_paragraph(
            "Considering your expertise in the subject matter, the Company is desirous of appointing your goodself as a trading advisor (hereinafter referred to as “TA”) with respect to its trading operations and undertake trading activities on its behalf including but not limited to trading, jobbing, arbitrage, hedging etc."
        )
        set_font(content, font_name="Roboto", font_size=11)
        doc.add_paragraph()  # Blank row

        # Terms and Conditions section
        doc.add_paragraph("The terms and conditions for your appointment would be as under:", style='Heading 1')

        terms_and_conditions = [
            "The Company shall provide all the necessary infrastructure to you, including trading terminals of the Bombay Stock Exchange and the National Stock Exchange, trading/algo software, charting software, Wi-Fi, internet, web access, television, furniture and fixtures, electricity, water, air conditioning, telephone/s lines, etc.",
            "We understand that you would be assisted by your team members (hereinafter referred to as 'the team/team members'), the details of which are as per Annexure 'A'.",
            "You and your team shall be permitted to use the facilities described in para 1 above. You and team shall devote your skill, time, ability and attention to conducting trading/jobbing/arbitrage/hedging transactions/strategies on the exchange/s in the best interest of the Company. You and your team shall be responsible for all trading-related activities of the team.",
            "You shall execute/instruct to execute all the transactions in the Unique Client Code of the Company and only for the Company. The Company shall also provide the requisite funds to enable you to undertake transactions on the exchange/s through a SEBI registered Broker in the Company’s client code.",
            "You and your team agree to keep an interest-free security deposit, as may be mutually agreed from time to time, the proceeds of which will be utilized by the Company at its sole discretion.",
            "You and your team agree to carry on said business in complete confidentiality and shall not divulge trading positions, strategies, etc., or any other information related to the said business to anyone whatsoever.",
            "You shall charge fees for the services rendered by you and your team and shall periodically provide necessary instructions for disbursal of the same. The Company will make payment to you and your team as detailed in the instruction note after appropriate deduction of tax at source as per the provisions of the Income-tax Act, 1961.",
            "As mentioned earlier, you may also, if agreed upon by the Company, appoint/employ other persons referred to as your team to assist you in your trading activity. Your team agrees to abide by the terms and conditions as set out in this letter and confirm the same. New team members may be introduced from time to time at your sole discretion. Necessary intimation will be sent by you to the Company on introduction or removal of any team member.",
            "You and your team shall implement the terms of this agreement in good faith and due diligence in the best interest of the Company.",
            "Nothing in this arrangement shall constitute or be deemed to constitute a partnership, relationship of principal and agent or employer and employee between any of the parties, and none of them shall have any authority to bind any of the other parties in any way except for the purposes of the business of the Company.",
            "You or your team members shall not undertake any personal trading in the Unique Client Code of the Company under any circumstances. This shall be construed as 'unauthorized' trading activity and liable for damages by the Company.",
            "You and your team shall hereby comply with all the applicable rules and regulations, without limitation to, SEBI (Prohibition of Fraudulent and Unfair Trade Practices Relating to Securities Market) Regulations, Exchange guidelines and regulations, and any amendments and changes thereto, or any other act/s, and any such guidelines as may be prescribed from time to time by the Company.",
            "This arrangement shall be for one year from April 01, 2018, but the Company reserves the right to terminate at any time without giving any prior notice or modify any terms and conditions of this letter from time to time as may be deemed necessary by the Company."
        ]

        for term in terms_and_conditions:
            paragraph = doc.add_paragraph(term, style='List Number')
            set_font(paragraph, font_name="Roboto", font_size=11)
            doc.add_paragraph()  # Blank row between terms

        # Signature section
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.style = 'Table Grid'
        sign_table.autofit = True

        sign_table.cell(0, 0).text = "For Elixir Equities Pvt Ltd,\n\n\n\nDipan Mehta\nDirector"
        sign_table.cell(0, 1).text = f"I Accept\n\n\n\nName: {current_trading_advisor} \nPAN : {current_trading_advisor_new['PAN'].iloc[0]}"

        set_table_font(sign_table, font_name="Roboto", font_size=11)

        # Annexure A - Start new page
        doc.add_page_break()
        doc.add_paragraph("Annexure - A", style='Heading 1')
        doc.add_paragraph(
            "The following is the list of team members of the TA presently working with him which may change from time to time."
        )

        # Get PAYEE list and PAN list
        current_payee = ta_payee_mapping[ta_payee_mapping['TRADING ADVISOR'] == current_trading_advisor]['PAYEE'].iloc[0]
        current_pan = ta_payee_mapping[ta_payee_mapping['TRADING ADVISOR'] == current_trading_advisor]['PAN'].iloc[0]

        table = doc.add_table(rows=len(current_payee)+1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Payee Name"
        table.cell(0, 1).text = "PAN"

        for i, (payee, pan) in enumerate(zip(current_payee, current_pan), start=1):
            table.cell(i, 0).text = payee
            table.cell(i, 1).text = pan

        set_table_font(table, font_name="Roboto", font_size=11)
        
        os.makedirs(f"EL-{year_folder}", exist_ok=True)
        doc.save(f"EL-{year_folder}/{current_trading_advisor}.docx")
