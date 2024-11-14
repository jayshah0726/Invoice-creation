from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

file_path = 'EW - Details of Professional Fees Paid for last 7 years.xlsx'
address_file_path = 'EW Master 1.csv' 
sheet_names = ['2017-18']

address_data = pd.read_csv(address_file_path)

def set_font(paragraph, font_name="Roboto", font_size=12):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

        # This is needed to ensure font name works with Word documents
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

# Function to set font for all text in a table
def set_table_font(table, font_name="Arial", font_size=12):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_font(paragraph, font_name=font_name, font_size=font_size)

for sheet in sheet_names:
    year_folder = sheet.replace('/', '-')
    # names and address - load from EW details of file
    master_data = pd.read_csv('EW Master 1.csv')
    data = pd.read_excel('EW - Details of Professional Fees Paid for last 7 years 1.xlsx', sheet_name=sheet)
    trading_advisor_list = list(data['TRADING ADVISOR'].drop_duplicates().dropna())
    filtered_data = data[['TRADING ADVISOR','PAYEE','PAN']]
    grouped_data_by_advisor = filtered_data.groupby('TRADING ADVISOR')
    ta_payee_mapping = grouped_data_by_advisor.agg({
        "PAN": list,
        "PAYEE": list
    }).reset_index()

    for current_trading_advisor in trading_advisor_list:

        # for ele in master_data['Name']:
        #     print(ele)
        #     print(len(ele))

        # break
        current_trading_advisor = current_trading_advisor.strip()
        print(current_trading_advisor)
        if current_trading_advisor == "MITESH DOSHI":
            current_trading_advisor = current_trading_advisor.replace(current_trading_advisor,"MITESH JAYANTIBHAI DOSHI")
        elif current_trading_advisor == "MITUL MORABIA":
            current_trading_advisor = current_trading_advisor.replace(current_trading_advisor,"MITUL MOHANLAL MORABIYA")
        # find the address from master data for the current trading advisor
        current_trading_advisor_new = master_data[(master_data['Name'] == current_trading_advisor) & (master_data['Year'] == '2017-18')]
        current_trading_advisor_address = current_trading_advisor_new['Address'].iloc[0]
        print(current_trading_advisor_address)
        
        # Initialize the document
        doc = Document()

        # Title
        title = doc.add_paragraph("ENGAGEMENT LETTER", style='Title')
        title.alignment = 1
        set_font(title,font_name="Roboto",font_size=14)

        # Date
        date = doc.add_paragraph(f"April 01, {sheet.split('-')[0]}", style='Normal')
        date.alignment = 0
        set_font(date,font_name="Roboto",font_size=12)

        # Add To section
        to_para = doc.add_paragraph()

        to_para.add_run(f"To,\n{current_trading_advisor},\n{current_trading_advisor_address}\n")
        set_font(to_para, font_name="Roboto", font_size=12)

        # Subject line
        subject_para = doc.add_paragraph()
        subject_para.add_run("Subject: Appointment as Trading Advisor\n\n").bold = True
        set_font(subject_para, font_name="Roboto", font_size=12)

        # Introduction
        intro = doc.add_paragraph(
            "Elixir Wealth Management Private Limited (hereinafter referred to as “the Company”) is a company incorporated under the erstwhile provisions of the Companies Act, 1956 carrying on the business of securities trading across various market segments."
        )
        set_font(intro, font_name="Roboto", font_size=12)

        # Content paragraphs
        content = doc.add_paragraph(
            "Considering your expertise in the subject matter, the Company is desirous of appointing your goodself as a trading advisor (hereinafter referred to as “TA”) with respect to its trading operations and undertake trading activities on its behalf including but not limited to trading, jobbing, arbitrage, hedging etc."
        )
        set_font(content, font_name="Roboto", font_size=12)

        # Adding "Terms and Conditions" title
        doc.add_paragraph("The terms and conditions for your appointment would be as under: \n", style='Heading 1')

        # List of terms and conditions as numbered list
        terms_and_conditions = [
            "The Company shall provide all the necessary infrastructure to you, including trading terminals of the Bombay Stock Exchange and the National Stock Exchange, trading/algo software, charting software, Wi-Fi, internet, web access, television, furniture and fixtures, electricity, water, air conditioning, telephone/s lines, etc.",
            "We understand that you would be assisted by your team members (hereinafter referred to as 'the team/team members'), the details of which are as per Annexure 'A'.",
            "You and your team shall be permitted to use the facilities described in para 1 above. You and team shall devote your skill, time, ability and attention to conducting trading/jobbing/arbitrage/hedging transactions/strategies on the exchange/s in the best interest of the Company. You and your team shall be responsible for all trading-related activities of the team.",
            "You shall execute/instruct to execute all the transactions in the Unique Client Code of the Company and only for the Company. The Company shall also provide the requisite funds to enable you to undertake transactions on the exchange/s through a SEBI registered Broker in the Company’s client code.",
            "You and your team agree to keep an interest-free security deposit, as may be mutually agreed from time to time, the proceeds of which will be utilized by the Company at its sole discretion.",
            "You and your team agree to carry on said business in complete confidentiality and shall not divulge trading positions, strategies, etc., or any other information related to the said business to anyone whatsoever.",
            "You shall charge fees for the services rendered by you and your team and shall periodically provide necessary instructions for disbursal of the same. The Company will make payment to you and your team as detailed in the instruction note after appropriate deduction of tax at source as per the provisions of the Income-tax Act, 1961.",
            "As mentioned earlier, you may also, if agreed upon by the Company, appoint/employ other persons referred to as your team to assist you in your trading activity. Your team agrees to abide by the terms and conditions as set out in this letter and confirm the same. New team members may be introduced from time to time at your sole discretion. Necessary intimation will be sent by you to the Company on introduction or removal of any team member.",
            "You and your team shall implement the terms of this agreement in good faith and due diligence in the best interest of the Company.",
            "Nothing in this arrangement shall constitute or be deemed to constitute a partnership, relationship of principal and agent or employer and employee between any of the parties, and none of them shall have any authority to bind any of the other parties in any way except for the purposes of the business of the Company.",
            "You or your team members shall not undertake any personal trading in the Unique Client Code of the Company under any circumstances. This shall be construed as 'unauthorized' trading activity and liable for damages by the Company.",
            "You and your team shall hereby comply with all the applicable rules and regulations, without limitation to, SEBI (Prohibition of Fraudulent and Unfair Trade Practices Relating to Securities Market) Regulations, Exchange guidelines and regulations, and any amendments and changes thereto, or any other act/s, and any such guidelines as may be prescribed from time to time by the Company.",
            "This arrangement shall be for one year from April 01, 2017, but the Company reserves the right to terminate at any time without giving any prior notice or modify any terms and conditions of this letter from time to time as may be deemed necessary by the Company."
        ]

        # Add each term as a numbered point
        for term in terms_and_conditions:
            paragraph = doc.add_paragraph(term, style='List Number')
            set_font(paragraph, font_name="Roboto", font_size=12)

        # Signature section
        sign_table = doc.add_table(rows=1, cols=2)
        sign_table.style = 'Table Grid'
        sign_table.autofit = True

        sign_table.cell(0, 0).text = "For Elixir Wealth Management Pvt Ltd,\n\n\n\nDipan Mehta\nDirector"
        sign_table.cell(0, 1).text = f"I Accept\n\n\n\nName: Dipan Mehta \nPAN : FFCPM9766K "

        set_table_font(sign_table,font_name="Roboto",font_size=12)

        # Annexure A
        doc.add_paragraph("\n\nAnnexure A", style='Heading 1')
        members = doc.add_paragraph(
            "The following is the list of team members of the TA presently working with him which may change from time to time as per commercial prudence of the TA."
        )
        set_font(members, font_name="Roboto", font_size=12)

        # table in the end
        filtered_ta_payee_mapping = ta_payee_mapping[ta_payee_mapping['TRADING ADVISOR'] == current_trading_advisor]

        for index,each_row in filtered_ta_payee_mapping.iterrows():
            trading_advisor = each_row['TRADING ADVISOR']
            pan_list = each_row["PAN"]
            payee_list = each_row['PAYEE']
            pan_payee_dict = {}
            for i in range(0,len(pan_list)):
                if payee_list[i] not in pan_payee_dict.keys() and payee_list[i] != trading_advisor:
                    pan_payee_dict[payee_list[i]] = pan_list[i]

            # insert these as table now
            # append the table in the end
            if len(pan_payee_dict.keys()) > 0:
                table = doc.add_table(rows=1,cols=3)
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'PAYEE'
                hdr_cells[1].text = 'PAN'
                hdr_cells[2].text = 'SIGNATURE'
                for key, value in pan_payee_dict.items():
                    modified_row = table.add_row().cells
                    modified_row[0].text = str(key)
                    modified_row[1].text = str(value)

                set_table_font(table,font_name="Roboto",font_size=12)

        # Save the document
        doc.save(f"EL(2017-18)/Engagement_Letter_{trading_advisor}.docx")